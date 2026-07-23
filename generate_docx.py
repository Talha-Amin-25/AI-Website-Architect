import os
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding (margins) in dxas (1/20 of a pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_with_spacing(doc, text, level, before=12, after=6):
    """Adds a heading with explicit spacing and color."""
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(before)
    heading.paragraph_format.space_after = Pt(after)
    heading.paragraph_format.keep_with_next = True
    
    # Custom colors
    run = heading.runs[0]
    run.font.name = 'Outfit'
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(15, 23, 42) # Slate-900 (#0F172A)
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(197, 160, 89) # Gold (#C5A059)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(71, 85, 105) # Slate-600
        
    return heading

def main():
    doc = Document()
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Set default style fonts
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Inter'
    font.size = Pt(11)
    font.color.rgb = RGBColor(31, 41, 55) # Gray-800 (#1F2937)
    
    # ==========================================
    # COVER PAGE
    # ==========================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(120)
    title_p.paragraph_format.space_after = Pt(8)
    
    title_run = title_p.add_run("AI WEBSITE ARCHITECT ENGINE")
    title_run.font.name = 'Outfit'
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 23, 42) # Slate-900
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(240)
    
    sub_run = subtitle_p.add_run("Dynamic Design Intelligence System & UX Blueprint Workspace")
    sub_run.font.name = 'Outfit'
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(197, 160, 89) # Gold
    
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(6)
    
    meta_run = meta_p.add_run("TECHNICAL DOCUMENTATION & USER MANUAL\n")
    meta_run.font.bold = True
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(107, 114, 128)
    
    date_run = meta_p.add_run(f"Version 2.5  |  Date: {datetime.date.today().strftime('%B %d, %Y')}\nTarget: Technical Evaluation Portfolio")
    date_run.font.size = Pt(9.5)
    date_run.font.color.rgb = RGBColor(156, 163, 175)
    
    doc.add_page_break()
    
    # ==========================================
    # 1. INTRODUCTION & OBJECTIVES
    # ==========================================
    add_heading_with_spacing(doc, "1. Executive Introduction", level=1, before=18)
    
    p = doc.add_paragraph(
        "The AI Website Architect project represents a shift from raw template duplication to algorithmic layout mapping. "
        "Unlike generic CMS builders that generate static content or populate pre-made themes, this system operates as a "
        "Senior UX Strategist and Information Architect. By processing client questionnaires, service packages, and image asset "
        "intelligence (composition, light ratios, and background tones), the engine outputs a Pydantic-validated design system "
        "and logical site blueprint. It answers not just 'what' content is presented, but 'where', 'why', and 'how' it must "
        "be structured to align with target audiences."
    )
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(10)
    
    # Highlight Box
    callout = doc.add_paragraph()
    callout.paragraph_format.left_indent = Inches(0.4)
    callout.paragraph_format.right_indent = Inches(0.4)
    callout.paragraph_format.space_before = Pt(8)
    callout.paragraph_format.space_after = Pt(12)
    
    run_callout = callout.add_run(
        "Core Paradigm: The core strategy is focused on logical information architecture. The program computes "
        "spatial layouts, visual focus alignments, WCAG 2.1 compliance parameters, and SEO indicators "
        "to formulate an optimal user conversion journey."
    )
    run_callout.font.italic = True
    run_callout.font.size = Pt(10)
    run_callout.font.color.rgb = RGBColor(75, 85, 99)
    
    # ==========================================
    # 2. SYSTEM ARCHITECTURE
    # ==========================================
    add_heading_with_spacing(doc, "2. System Architecture & Algorithms", level=1)
    
    p = doc.add_paragraph(
        "The architecture is modularly separated to facilitate ease of deployment, updates, and maintenance. "
        "The system consists of two primary layers:"
    )
    p.paragraph_format.space_after = Pt(8)
    
    bullets = [
        ("FastAPI Application Server (app.py)", "Serves RESTful API endpoints and serves the compiled workspace interface from the static folder."),
        ("Pydantic Schema Validation (engine/models.py)", "Enforces strict model typing, sanitizing hex inputs, score limits, and output schemas."),
        ("Design Intelligence Engine (engine/architect.py)", "Executes relative luminance checking, weighted hero selection, and section priority ordering."),
        ("Design Tokens Templates (engine/templates.py)", "Stores font pairs, default layouts, and SEO/mobile checkpoints.")
    ]
    
    for title, desc in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        run_title = bp.add_run(f"{title}: ")
        run_title.bold = True
        bp.add_run(desc)

    add_heading_with_spacing(doc, "2.1 Hero Selection Formula", level=2)
    
    p = doc.add_paragraph(
        "Rather than picking visual media randomly, the engine selects the primary hero visual using a weighted "
        "attribute equation from image intelligence scores:"
    )
    p.paragraph_format.space_after = Pt(8)
    
    formula_p = doc.add_paragraph()
    formula_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_p.paragraph_format.space_before = Pt(6)
    formula_p.paragraph_format.space_after = Pt(6)
    f_run = formula_p.add_run("Score = (0.35 * Lighting) + (0.35 * Composition) + (0.20 * Luxury) + (0.10 * Contrast)")
    f_run.font.bold = True
    f_run.font.name = 'Consolas'
    f_run.font.size = Pt(10.5)
    f_run.font.color.rgb = RGBColor(197, 160, 89)
    
    p = doc.add_paragraph(
        "This formula ranks assets by balancing clarity, artistic structure, brand feel, and the availability "
        "of negative space for overlay text."
    )
    p.paragraph_format.space_after = Pt(10)
    
    # ==========================================
    # 3. SETUP & HOW TO RUN
    # ==========================================
    add_heading_with_spacing(doc, "3. Installation & Setup Guide", level=1)
    
    p = doc.add_paragraph(
        "To ensure this application runs smoothly on any device (Windows, macOS, Linux, or server environments), "
        "it utilizes basic Python packages and serves static assets locally. Follow these universal steps:"
    )
    p.paragraph_format.space_after = Pt(10)
    
    steps = [
        ("Step 1: Check Python version", "Open a terminal (Command Prompt, PowerShell, or Terminal on macOS/Linux) and verify that Python 3.8+ is installed:\npython --version"),
        ("Step 2: Install dependencies", "Navigate to the project root and install the required modules using the requirements.txt file:\npip install -r requirements.txt\n(Note: On Windows systems with multiple environments, you can use: py -m pip install -r requirements.txt)"),
        ("Step 3: Run the server", "Start the FastAPI application by running the entry point script:\npython app.py\n(Alternatively: py app.py)"),
        ("Step 4: Launch the dashboard", "Once the logs indicate 'Application startup complete', open your web browser and go to:\nhttp://localhost:8000")
    ]
    
    for title, cmd in steps:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(6)
        r_title = sp.add_run(f"{title}\n")
        r_title.bold = True
        r_title.font.size = Pt(11.5)
        
        r_cmd = sp.add_run(cmd)
        r_cmd.font.name = 'Consolas'
        r_cmd.font.size = Pt(9.5)
        r_cmd.font.color.rgb = RGBColor(79, 70, 229)
        
    add_heading_with_spacing(doc, "3.1 Multi-Device Troubleshooting", level=2)
    
    p = doc.add_paragraph(
        "Port Conflict: If port 8000 is occupied on the target machine, edit the bottom lines in app.py to configure a different port (e.g. port=8080):\n"
        "uvicorn.run('app:app', host='0.0.0.0', port=8080)\n\n"
        "No Module Pip: If python is an embedded version and pip is missing, download the bootstrap script 'get-pip.py' from bootstrap.pypa.io and run it using your python executable."
    )
    p.runs[0].font.name = 'Consolas'
    p.runs[0].font.size = Pt(9.5)
    p.paragraph_format.space_after = Pt(12)

    # ==========================================
    # 4. HOW TO USE THE DASHBOARD
    # ==========================================
    add_heading_with_spacing(doc, "4. Interactive Dashboard User Guide", level=1)
    
    p = doc.add_paragraph(
        "The workspace is divided into three key panels which reflect updates in real-time:"
    )
    p.paragraph_format.space_after = Pt(8)
    
    panels = [
        ("Left Navigation Sidebar", "Includes navigation tabs and a dynamic progress ring. The circular progress indicator animates to show the overall consistency score of the project (calculated from layout, UX, and speed metrics)."),
        ("Center Interactive Workspace", "Contains the workspace tabs:\n"
         "- Blueprint: Displays connected layout nodes with design decisions and confidence scores.\n"
         "- Wireframe: Renders visual layout skeletons using the extracted primary, secondary, and text colors.\n"
         "- SEO & Accessibility: Displays priority checklist recommendations.\n"
         "- JSON Output: Presents the exact raw JSON output to copy or export directly."),
        ("Right Specifications Panel", "Highlights Asset Intelligence scores (lighting, color harmony, typography pairing status), typographic font sizes, color palette swatch hexes with accessibility ratings, and a Performance Metrics line chart.")
    ]
    
    for name, desc in panels:
        p_card = doc.add_paragraph(style='List Bullet')
        p_card.paragraph_format.space_after = Pt(6)
        r_n = p_card.add_run(f"{name}: ")
        r_n.bold = True
        p_card.add_run(desc)

    add_heading_with_spacing(doc, "4.1 Generating a New Design Blueprint", level=2)
    
    p = doc.add_paragraph(
        "1. Click the 'New Blueprint' button on the left sidebar navigation menu.\n"
        "2. An input form modal will slide into view pre-filled with example luxury hospitality values.\n"
        "3. Modify fields such as brand hex codes, business category (e.g. hospitality, software startup, law firm), service packages, and image asset scores.\n"
        "4. Click 'Run UX Intelligence Engine'. The dashboard will retrieve the generated blueprint, and the UI panels will animate to showcase the updated design style."
    )
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(12)

    # Save Document
    doc.save("AI_Website_Architect_Documentation.docx")
    print("AI_Website_Architect_Documentation.docx generated successfully.")

if __name__ == "__main__":
    main()
